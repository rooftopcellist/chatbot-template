"""
Document processor module for loading and processing documents.
"""

import os
import json
import frontmatter
from typing import List
from llama_index.core.schema import Document
import config

# Import specialized document processing libraries
import docx
import pandas as pd
from pypdf import PdfReader
import asciidoc3

def load_documents() -> List[Document]:
    """
    Load all supported documents from the configured directory.
    Supports multiple file types: .md, .docx, .pdf, .csv, .json, .log, .adoc, .rst

    Returns:
        List[Document]: List of processed documents
    """
    documents = []

    # Ensure the docs directory exists
    if not os.path.exists(config.DOCS_DIR):
        print(f"Warning: Documents directory '{config.DOCS_DIR}' not found.")
        return documents

    # File extension to processor function mapping
    processors = {
        '.md': process_markdown_file,
        '.docx': process_docx_file,
        '.pdf': process_pdf_file,
        '.csv': process_csv_file,
        '.json': process_json_file,
        '.log': process_log_file,
        '.adoc': process_asciidoc_file,
        '.rst': process_rst_file,
    }

    # Verify that all supported extensions have processors
    for ext in config.SUPPORTED_EXTENSIONS:
        if ext not in processors:
            print(f"Warning: Supported extension '{ext}' has no processor defined.")

    # Walk through the directory and process all supported files
    for root, _, files in os.walk(config.DOCS_DIR):
        for file in files:
            file_path = os.path.join(root, file)
            _, ext = os.path.splitext(file)

            if ext.lower() in processors:
                try:
                    # Process the file with the appropriate processor
                    doc = processors[ext.lower()](file_path)
                    if doc:
                        documents.append(doc)
                except Exception as e:
                    print(f"Error processing file {file_path}: {e}")

    print(f"Loaded {len(documents)} documents from {config.DOCS_DIR}")
    return documents

def process_markdown_file(file_path: str) -> Document:
    """
    Process a single Markdown file.
    Strips YAML front matter if present.

    Args:
        file_path (str): Path to the Markdown file

    Returns:
        Document: Processed document
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        # Parse frontmatter
        post = frontmatter.load(f)

        # Extract content without frontmatter
        content = post.content

        # Create metadata
        metadata = {
            'source': file_path,
            'filename': os.path.basename(file_path),
            'filetype': 'markdown',
            **post.metadata  # Include frontmatter as metadata
        }

        # Create Document object
        return Document(text=content, metadata=metadata)

def process_docx_file(file_path: str) -> Document:
    """
    Process a single DOCX file.

    Args:
        file_path (str): Path to the DOCX file

    Returns:
        Document: Processed document
    """
    doc = docx.Document(file_path)

    # Extract text from paragraphs
    content = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    # Create metadata
    metadata = {
        'source': file_path,
        'filename': os.path.basename(file_path),
        'filetype': 'docx',
    }

    # Create Document object
    return Document(text=content, metadata=metadata)

def process_pdf_file(file_path: str) -> Document:
    """
    Process a single PDF file.

    Args:
        file_path (str): Path to the PDF file

    Returns:
        Document: Processed document
    """
    reader = PdfReader(file_path)

    # Extract text from pages
    content = ""
    for page in reader.pages:
        content += page.extract_text() + "\n\n"

    # Create metadata
    metadata = {
        'source': file_path,
        'filename': os.path.basename(file_path),
        'filetype': 'pdf',
        'pages': len(reader.pages),
    }

    # Create Document object
    return Document(text=content, metadata=metadata)

def process_csv_file(file_path: str) -> Document:
    """
    Process a single CSV file.

    Args:
        file_path (str): Path to the CSV file

    Returns:
        Document: Processed document
    """
    # Read CSV file with pandas
    df = pd.read_csv(file_path)

    # Convert to string representation
    content = df.to_string()

    # Create metadata
    metadata = {
        'source': file_path,
        'filename': os.path.basename(file_path),
        'filetype': 'csv',
        'rows': len(df),
        'columns': len(df.columns),
        'column_names': df.columns.tolist(),
    }

    # Create Document object
    return Document(text=content, metadata=metadata)

def process_json_file(file_path: str) -> Document:
    """
    Process a single JSON file.

    Args:
        file_path (str): Path to the JSON file

    Returns:
        Document: Processed document
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        # Load JSON
        data = json.load(f)

        # Convert to formatted string
        content = json.dumps(data, indent=2)

    # Create metadata
    metadata = {
        'source': file_path,
        'filename': os.path.basename(file_path),
        'filetype': 'json',
    }

    # Create Document object
    return Document(text=content, metadata=metadata)

def process_log_file(file_path: str) -> Document:
    """
    Process a single log file.

    Args:
        file_path (str): Path to the log file

    Returns:
        Document: Processed document
    """
    with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
        content = f.read()

    # Create metadata
    metadata = {
        'source': file_path,
        'filename': os.path.basename(file_path),
        'filetype': 'log',
    }

    # Create Document object
    return Document(text=content, metadata=metadata)

def process_asciidoc_file(file_path: str) -> Document:
    """
    Process a single AsciiDoc file.

    Args:
        file_path (str): Path to the AsciiDoc file

    Returns:
        Document: Processed document
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Convert AsciiDoc to text (if needed)
    # Note: asciidoc3 is imported but we're just using the raw content for now
    # This could be enhanced to use asciidoc3 for proper conversion

    # Create metadata
    metadata = {
        'source': file_path,
        'filename': os.path.basename(file_path),
        'filetype': 'asciidoc',
    }

    # Create Document object
    return Document(text=content, metadata=metadata)

def process_rst_file(file_path: str) -> Document:
    """
    Process a single reStructuredText (.rst) file.
    Converts RST to plain text while preserving structure.

    Args:
        file_path (str): Path to the RST file

    Returns:
        Document: Processed document
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        rst_content = f.read()

    # For RST files, we'll use the raw content directly for better RAG performance
    # This avoids parsing issues with undefined substitutions and directives
    # while still preserving the readable structure of RST
    content = rst_content
    title = ''

    # Try to extract title from RST content using simple parsing
    try:
        lines = rst_content.split('\n')
        for i, line in enumerate(lines):
            # Look for RST title patterns (underlined text)
            if line.strip() and i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                # Check if next line is all underline characters (=, -, *, etc.)
                if (len(next_line) >= len(line.strip()) and
                    next_line and
                    all(c in '=-*^~"\'`#' for c in next_line)):
                    title = line.strip()
                    break
    except Exception:
        # If title extraction fails, continue without title
        pass

    # Create metadata
    metadata = {
        'source': file_path,
        'filename': os.path.basename(file_path),
        'filetype': 'rst',
    }

    # Add title to metadata if available
    if title:
        metadata['title'] = title

    # Create Document object
    return Document(text=content, metadata=metadata)
